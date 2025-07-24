import os
import streamlit as st
import pdfplumber
from docx import Document as DocxFile

from langchain_community.vectorstores import FAISS
from langchain.docstore.document import Document
from langchain.chains import RetrievalQA
from langchain.prompts import PromptTemplate
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langgraph.graph import StateGraph
from langchain_core.runnables import RunnableLambda

from langchain_google_genai import ChatGoogleGenerativeAI, GoogleGenerativeAIEmbeddings
from langchain_community.tools import DuckDuckGoSearchRun
from dotenv import load_dotenv
load_dotenv()

# --- ENVIRONMENT & MODELS ---
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")
if not GOOGLE_API_KEY:
    raise RuntimeError("Missing GOOGLE_API_KEY in environment.")

llm_model = ChatGoogleGenerativeAI(model="gemini-1.5-flash", temperature=0.2)
embedder = GoogleGenerativeAIEmbeddings(model="models/embedding-001")
splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
web_search_tool = DuckDuckGoSearchRun()

# --- DOCUMENT INGESTION ---
def parse_file_content(path):
    if path.endswith(".pdf"):
        with pdfplumber.open(path) as pdf:
            return "\n".join([p.extract_text() for p in pdf.pages if p.extract_text()])
    elif path.endswith(".txt"):
        with open(path, "r", encoding="utf-8") as f:
            return f.read()
    elif path.endswith(".docx"):
        doc = DocxFile(path)
        return "\n".join([p.text for p in doc.paragraphs])
    return ""

def gather_documents_from_folder(folder):
    docs = []
    for fname in os.listdir(folder):
        if fname.lower().endswith(('.pdf', '.txt', '.docx')):
            text = parse_file_content(os.path.join(folder, fname))
            if text:
                docs.append(text)
    return docs

# --- AGENT DEFINITIONS ---
def decision_agent(state):
    user_question = state.get("question", "")
    prompt = PromptTemplate.from_template(
        "Given the following user query, decide which resource to use: [internet, local, direct].\nQuery: {question}\nYour answer:"
    )
    result = (prompt | llm_model).invoke({"question": user_question}).content.lower()
    if "internet" in result:
        path = "internet"
    elif "local" in result:
        path = "local"
    else:
        path = "direct"
    return {**state, "chosen_path": path}

def internet_agent(state):
    q = state["question"]
    try:
        web_result = web_search_tool.run(q)
        return {**state, "info": web_result}
    except Exception as err:
        return {**state, "info": f"Internet search error: {err}"}

def local_agent(state):
    q = state["question"]
    retriever = state["kb_retriever"]
    qa = RetrievalQA.from_chain_type(llm=llm_model, retriever=retriever)
    response = qa.run(q)
    return {**state, "info": response}

def direct_agent(state):
    q = state["question"]
    reply = llm_model.invoke(q)
    return {**state, "info": reply.content}

def synthesis_agent(state):
    info = state["info"]
    summary_prompt = PromptTemplate.from_template(
        "Please provide a concise, clear summary of the following information:\n\n{info}"
    )
    summary = (summary_prompt | llm_model).invoke({"info": info}).content
    return {**state, "final_answer": summary}

# --- WORKFLOW GRAPH ---
def execute_workflow(user_question, kb_retriever):
    graph = StateGraph(dict)
    graph.set_entry_point("decider")

    graph.add_node("decider", RunnableLambda(decision_agent))
    graph.add_node("internet", RunnableLambda(internet_agent))
    graph.add_node("local", RunnableLambda(local_agent))
    graph.add_node("direct", RunnableLambda(direct_agent))
    graph.add_node("synthesizer", RunnableLambda(synthesis_agent))

    def path_selector(state): return state["chosen_path"]
    graph.add_conditional_edges("decider", path_selector, {
        "internet": "internet",
        "local": "local",
        "direct": "direct"
    })

    for node in ["internet", "local", "direct"]:
        graph.add_edge(node, "synthesizer")

    graph.set_finish_point("synthesizer")
    compiled = graph.compile()
    return compiled.invoke({"question": user_question, "kb_retriever": kb_retriever})["final_answer"]

# --- STREAMLIT UI ---
st.set_page_config(page_title="🔎 Research Copilot", layout="centered")
st.title("🤖 Next-Gen Agentic Researcher")

kb_retriever = None
has_docs = False

if os.path.exists("rag"):
    with st.spinner("🔄 Scanning 'rag' directory for documents..."):
        docs = gather_documents_from_folder("rag")
        if docs:
            doc_chunks = splitter.create_documents(docs)
            store = FAISS.from_documents(doc_chunks, embedder)
            kb_retriever = store.as_retriever()
            has_docs = True
            st.success(f"Loaded {len(docs)} document(s).")
        else:
            st.warning("No supported files found in 'rag'.")

if not has_docs:
    st.info("Using default fallback knowledge base.")
    fallback_docs = [
        Document(page_content="LangGraph is a Python framework for agentic workflows."),
        Document(page_content="Gemini 1.5 Flash excels at fast summarization tasks."),
    ]
    store = FAISS.from_documents(fallback_docs, embedder)
    kb_retriever = store.as_retriever()

user_input = st.text_input("Type your research question:", placeholder="e.g. How does LangGraph work?")
ask = st.button("Ask")

if ask:
    if not user_input.strip():
        st.warning("Please enter a valid question.")
    else:
        with st.spinner("🧠 Processing..."):
            try:
                result = execute_workflow(user_input, kb_retriever)
                st.success("Complete!")
                st.subheader("📙 Result:")
                st.write(result)
            except Exception as e:
                st.error(f"❌ Error: {str(e)}")